"""
Render docs/PROPOSAL.md to an editable Word document.

    python tools/render_proposal_docx.py     # writes docs/KVStream-Proposal-rev2.docx

Requires `pip install python-docx`. The .docx is a build product: edit
docs/PROPOSAL.md and re-run this. If you edit the Word file directly — which is
the point of having one — that copy becomes the source and this script will
overwrite it, so save it under a new name first.

Everything is applied through **built-in Word styles** (Heading 1-4, Quote,
List Bullet, List Number, Table Grid) rather than direct formatting, so the
whole document can be restyled from Word's style pane without touching a single
paragraph. Only three things are set directly: the theme fonts, the heading
colour, and a monospace character style for identifiers, because Word has no
built-in code style.
"""

from __future__ import annotations

import pathlib
import re
import sys

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = pathlib.Path(sys.argv[1] if len(sys.argv) > 1 else ".")
SOURCE = ROOT / "docs" / "PROPOSAL.md"
OUT = ROOT / "docs" / "KVStream-Proposal-rev2.docx"

INK = RGBColor(0x16, 0x18, 0x1A)
ACCENT = RGBColor(0x2D, 0x4A, 0x7C)
MUTED = RGBColor(0x5A, 0x61, 0x69)

BODY_FONT = "Georgia"
HEAD_FONT = "Segoe UI"
MONO_FONT = "Consolas"

# **bold** | *italic* | `code` | [text](url) — anything else is plain text.
INLINE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*)|`[^`]+`|\[[^\]]+\]\([^)]+\))")


# ----------------------------------------------------------------------
# document setup
# ----------------------------------------------------------------------


def build_document() -> Document:
    doc = Document()

    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)   # A4
    section.top_margin = section.bottom_margin = Cm(2.0)
    section.left_margin = section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.12
    _set_east_asian(normal, BODY_FONT)

    for level, (size, space_before) in enumerate(
        [(20, 0), (14, 16), (11.5, 12), (10.5, 10)], start=1
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = HEAD_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.italic = False
        style.font.color.rgb = INK if level == 1 else ACCENT
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True
        _set_east_asian(style, HEAD_FONT)

    quote = doc.styles["Quote"]
    quote.font.name = BODY_FONT
    quote.font.size = Pt(10)
    quote.font.italic = False
    quote.font.color.rgb = INK
    quote.paragraph_format.left_indent = Cm(0.6)
    quote.paragraph_format.space_after = Pt(6)

    # Word ships no code style; identifiers need one and it must be a real
    # style so it can be restyled rather than hunted down run by run.
    code = doc.styles.add_style("Code Inline", WD_STYLE_TYPE.CHARACTER)
    code.font.name = MONO_FONT
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor(0x1F, 0x35, 0x57)

    _add_page_numbers(section)
    return doc


def _set_east_asian(style, font_name: str) -> None:
    """python-docx sets only the Latin face; Word also reads eastAsia/cs."""
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:eastAsia", "w:cs", "w:hAnsi"):
        rfonts.set(qn(attr), font_name)


def _add_page_numbers(section) -> None:
    para = section.footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    for kind, text in (("begin", None), (None, "PAGE"), ("end", None)):
        if kind:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), kind)
        else:
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = text
        run._r.append(el)
    run.font.name = HEAD_FONT
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


# ----------------------------------------------------------------------
# inline runs
# ----------------------------------------------------------------------


def write_inline(paragraph, text: str, *, bold: bool = False, italic: bool = False) -> None:
    """
    Emit `text` into `paragraph`, honouring markdown inline markup.

    Recursive, because emphasis nests: `**a `code` b**` is common in this
    document and a single pass would match the outer bold and emit its contents
    flat, leaving literal backticks in the Word file.
    """
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            write_inline(paragraph, part[2:-2], bold=True, italic=italic)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1], style="Code Inline")
            if bold:
                run.bold = True
        elif part.startswith("[") and "](" in part:
            run = paragraph.add_run(part[1 : part.index("]")])
            run.font.color.rgb = ACCENT
            run.underline = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            write_inline(paragraph, part[1:-1], bold=bold, italic=True)
        else:
            run = paragraph.add_run(part.replace("<br>", ""))
            if bold:
                run.bold = True
            if italic:
                run.italic = True


# ----------------------------------------------------------------------
# block parsing
# ----------------------------------------------------------------------


def is_table_row(line: str) -> bool:
    return line.startswith("|") and line.endswith("|")


def split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def render(doc: Document, lines: list[str]) -> None:
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Horizontal rule — Word has no <hr>; an empty bottom-bordered
        # paragraph is the idiomatic equivalent and stays editable.
        if set(stripped) <= {"-", "*"} and len(stripped) >= 3 and not is_table_row(stripped):
            _horizontal_rule(doc)
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            para = doc.add_paragraph(style=f"Heading {min(level, 4)}")
            write_inline(para, stripped.lstrip("#").strip())
            i += 1
            continue

        if stripped.startswith(">"):
            block, i = _gather(lines, i, lambda ln: ln.strip().startswith(">"))
            text = " ".join(ln.strip().lstrip(">").strip() for ln in block if ln.strip() != ">")
            for chunk in text.split("  "):
                if chunk.strip():
                    write_inline(doc.add_paragraph(style="Quote"), chunk.strip())
            i += 1 if i < len(lines) else 0
            continue

        if is_table_row(stripped):
            rows, i = _gather(lines, i, lambda ln: is_table_row(ln.strip()))
            _table(doc, [split_row(r) for r in rows])
            continue

        if re.match(r"^[-*]\s+", stripped):
            block, i = _gather(lines, i, lambda ln: bool(ln.strip()) and not ln.startswith("#"))
            for item in _list_items(block, r"^[-*]\s+"):
                write_inline(doc.add_paragraph(style="List Bullet"), item)
            continue

        if re.match(r"^\d+\.\s+", stripped):
            block, i = _gather(lines, i, lambda ln: bool(ln.strip()) and not ln.startswith("#"))
            for item in _list_items(block, r"^\d+\.\s+"):
                write_inline(doc.add_paragraph(style="List Number"), item)
            continue

        # Ordinary paragraph: markdown soft-wraps, so join until a blank line.
        block, i = _gather(
            lines, i,
            lambda ln: bool(ln.strip()) and not ln.strip().startswith(("#", ">", "|", "---")),
        )
        text = " ".join(ln.strip() for ln in block)
        for chunk in text.split("<br>"):
            if chunk.strip():
                write_inline(doc.add_paragraph(), chunk.strip())


def _gather(lines, start, predicate):
    """Consume consecutive lines matching `predicate`; return (block, next_index)."""
    out = []
    i = start
    while i < len(lines) and predicate(lines[i]):
        out.append(lines[i])
        i += 1
    return out, i


def _list_items(block: list[str], marker: str) -> list[str]:
    """Re-join continuation lines onto their bullet."""
    items: list[str] = []
    for line in block:
        text = line.strip()
        if re.match(marker, text):
            items.append(re.sub(marker, "", text))
        elif items:
            items[-1] += " " + text
    return items


def _horizontal_rule(doc: Document) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(10)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "C3C8C1")
    borders.append(bottom)
    para._p.get_or_add_pPr().append(borders)


def _table(doc: Document, rows: list[list[str]]) -> None:
    # Row 1 is the header, row 2 is the |---|---| separator.
    body = [r for r in rows[1:] if not all(set(c) <= {"-", ":"} for c in r if c)]
    header = rows[0]
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for col, text in enumerate(header):
        cell = table.rows[0].cells[col]
        cell.paragraphs[0].text = ""
        write_inline(cell.paragraphs[0], text, bold=True)
        _shade(cell, "EEF1F6")

    for r, row in enumerate(body, start=1):
        for col in range(len(header)):
            cell = table.rows[r].cells[col]
            cell.paragraphs[0].text = ""
            write_inline(cell.paragraphs[0], row[col] if col < len(row) else "")

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(9)
                    if run.style.name != "Code Inline":
                        run.font.name = HEAD_FONT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _shade(cell, hex_colour: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(shd)


def main() -> None:
    doc = build_document()
    render(doc, SOURCE.read_text(encoding="utf-8").split("\n"))
    doc.core_properties.title = "KVStream — A Concurrency Gateway for Microsoft Foundry Local"
    doc.core_properties.subject = "Technical Proposal, Revision 2"
    doc.core_properties.author = "Shreyan Fernandes"
    doc.save(OUT)
    print(f"wrote {OUT}  ({OUT.stat().st_size / 1024:.0f} KB)")


if __name__ == "__main__":
    main()
